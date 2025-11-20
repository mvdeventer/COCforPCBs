"""
COC Report Generator
Compares BOM versions and analyzes schematics/assembly drawings
Configurable for any product and version comparison
"""

import io
import json
import logging
import re
import sys
import tkinter as tk
from datetime import datetime
from pathlib import Path
from tkinter import scrolledtext

import fitz  # PyMuPDF
import pandas as pd
import PyPDF2
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

# ============================================================================
# DEBUG MODE CONFIGURATION
# ============================================================================
DEBUG_MODE = "--debug" in sys.argv or "-d" in sys.argv

# Configure logging
if DEBUG_MODE:
    logging.basicConfig(
        level=logging.DEBUG,
        format="%(asctime)s - %(levelname)s - [%(filename)s:%(lineno)d] - %(message)s",
        handlers=[
            logging.FileHandler("coc_debug.log"),
            logging.StreamHandler(sys.stdout),
        ],
    )
    print("\n" + "=" * 60)
    print("🐛 DEBUG MODE ENABLED")
    print("=" * 60)
    print("Log file: coc_debug.log")
    print(f"Command line args: {sys.argv}")
    print("=" * 60 + "\n")
else:
    logging.basicConfig(
        level=logging.INFO,
        format="%(message)s",
        handlers=[logging.StreamHandler(sys.stdout)],
    )

logger = logging.getLogger(__name__)

# ============================================================================
# CONFIGURATION - Auto-detected from files in workspace
# ============================================================================
CONFIG = {
    "product_name": "",
    "version_old": "",
    "version_new": "",
    "bom_old_filename": "",
    "bom_new_filename": "",
    "schematic_old_filename": "",
    "schematic_new_filename": "",
    "assembly_old_filename": "",
    "assembly_new_filename": "",
}


def extract_bom_metadata(bom_file_path):
    """
    Extract version, variant, and product info from BOM file contents.

    logger.debug(f"extract_bom_metadata: Processing {bom_file_path}")

    This function is CRITICAL and runs BEFORE any BOM analysis.
    It searches the first 15 rows of the Excel file for metadata fields:
    - Version: Product version number (e.g., "9", "10")
    - Variant: Product variant description (e.g., "MLM24 PSU2", "NO SWITCH")
    - Variant Letter: Single letter variant identifier (e.g., "A", "B")
    - Product Name: Full product name
    - Assembly No: Assembly number with embedded product code
    - Build: Build identifier

    The function searches flexibly across all columns to handle different
    BOM file layouts and column positions.

    Args:
        bom_file_path: Path to the Excel BOM file

    Returns:
        dict: Metadata dictionary with extracted values (None if not found)
    """
    import openpyxl

    metadata = {
        "version": None,
        "variant": None,
        "variant_letter": None,
        "product_name": None,
        "assembly_no": None,
        "build": None,
    }

    try:
        wb = openpyxl.load_workbook(bom_file_path)
        ws = wb.active

        # Read first 15 rows to find metadata
        for row in ws.iter_rows(max_row=15, values_only=True):
            if row:
                # Search across all columns for metadata fields (flexible column detection)
                for i in range(len(row) - 1):
                    key = str(row[i]).strip() if row[i] else ""
                    value = (
                        str(row[i + 1]).strip()
                        if i + 1 < len(row) and row[i + 1]
                        else ""
                    )

                    if not key or not value:
                        continue

                    if "Prod Name" in key:
                        metadata["product_name"] = value
                    elif "Assembly No" in key:
                        metadata["assembly_no"] = value
                        # Extract product from assembly number (e.g., BT3411A-8[1])
                        match = re.search(r"([A-Z0-9]+)-", value)
                        if match and not metadata["product_name"]:
                            metadata["product_name"] = match.group(1)
                    elif key == "Variant:":
                        if not metadata["variant"]:
                            metadata["variant"] = value
                        elif len(value) == 1:  # Single letter variant
                            metadata["variant_letter"] = value
                    elif key == "Build:":
                        metadata["build"] = value
                    elif key == "Version:":
                        metadata["version"] = value

        wb.close()
    except Exception as e:
        print(f"Warning: Could not extract metadata from {bom_file_path.name}: {e}")

    return metadata


def load_file_config(workspace):
    """Load saved file configuration"""
    config_file = workspace / "input_files_config.json"
    if config_file.exists():
        try:
            with open(config_file, "r") as f:
                return json.load(f)
        except Exception as e:
            logger.warning(f"Failed to load file config: {e}")
    return None


def save_file_config(workspace, config):
    """Save file configuration"""
    config_file = workspace / "input_files_config.json"
    config["last_updated"] = datetime.now().isoformat()
    try:
        with open(config_file, "w") as f:
            json.dump(config, f, indent=2)
        logger.debug(f"Saved file config to {config_file}")
        return True
    except Exception as e:
        logger.error(f"Failed to save file config: {e}")
        return False


def show_file_selector_dialog(workspace):
    """Show dialog to select input files"""
    import tkinter as tk
    from tkinter import filedialog, messagebox

    selected_files = {}

    root = tk.Tk()
    root.title("Input File Configuration")
    root.geometry("700x550")

    # Header
    header_frame = tk.Frame(root, bg="#2c3e50", height=60)
    header_frame.pack(fill="x")
    header_frame.pack_propagate(False)

    header_label = tk.Label(
        header_frame,
        text="Select Input Files for COC Report",
        font=("Arial", 14, "bold"),
        bg="#2c3e50",
        fg="white",
    )
    header_label.pack(pady=15)

    # Content frame
    content_frame = tk.Frame(root, padx=30, pady=20)
    content_frame.pack(fill="both", expand=True)

    # Load existing config
    existing_config = load_file_config(workspace)

    # File selection entries
    file_entries = {}
    file_labels = {
        "bom_old": "Old BOM (Excel):",
        "bom_new": "New BOM (Excel):",
        "schematic_old": "Old Schematic (PDF):",
        "schematic_new": "New Schematic (PDF):",
        "assembly_old": "Old Assembly Drawing (PDF):",
        "assembly_new": "New Assembly Drawing (PDF):",
    }

    search_dir = (
        workspace / "input_files" if (workspace / "input_files").exists() else workspace
    )

    def browse_file(key, file_type):
        if file_type == "excel":
            filetypes = [("Excel files", "*.xlsx"), ("All files", "*.*")]
        else:
            filetypes = [("PDF files", "*.PDF *.pdf"), ("All files", "*.*")]

        filename = filedialog.askopenfilename(
            initialdir=search_dir,
            title=f"Select {file_labels[key]}",
            filetypes=filetypes,
        )

        if filename:
            file_entries[key].delete(0, tk.END)
            file_entries[key].insert(0, filename)
            selected_files[key] = filename

    row = 0
    for key, label_text in file_labels.items():
        # Label
        label = tk.Label(content_frame, text=label_text, font=("Arial", 10, "bold"))
        label.grid(row=row, column=0, sticky="w", pady=8)

        # Entry
        entry = tk.Entry(content_frame, width=40, font=("Arial", 9))
        entry.grid(row=row, column=1, sticky="ew", pady=8, padx=(10, 5))

        # Pre-fill with existing config
        if existing_config and existing_config.get(key):
            entry.insert(0, existing_config[key])
            selected_files[key] = existing_config[key]

        file_entries[key] = entry

        # Browse button
        file_type = "excel" if "bom" in key else "pdf"
        browse_btn = tk.Button(
            content_frame,
            text="Browse...",
            command=lambda k=key, ft=file_type: browse_file(k, ft),
            width=10,
            font=("Arial", 9),
            bg="#3498db",
            fg="white",
        )
        browse_btn.grid(row=row, column=2, pady=8, padx=(5, 0))

        row += 1

    content_frame.columnconfigure(1, weight=1)

    # Info label
    info_label = tk.Label(
        content_frame,
        text="Files will be remembered for future runs",
        font=("Arial", 9, "italic"),
        fg="#7f8c8d",
    )
    info_label.grid(row=row, column=0, columnspan=3, pady=(15, 10))

    # Buttons
    button_frame = tk.Frame(content_frame)
    button_frame.grid(row=row + 1, column=0, columnspan=3, pady=20)

    result = [None]

    def save_and_continue():
        # Validate all files selected
        missing = [
            file_labels[k]
            for k in file_labels
            if k not in selected_files or not selected_files[k]
        ]
        if missing:
            messagebox.showwarning(
                "Missing Files",
                f"Please select all files:\n" + "\n".join(f"- {m}" for m in missing),
            )
            return

        # Validate files exist
        for key, path in selected_files.items():
            if not Path(path).exists():
                messagebox.showerror("File Not Found", f"File does not exist:\n{path}")
                return

        # Save configuration
        if save_file_config(workspace, selected_files):
            result[0] = selected_files
            root.destroy()
        else:
            messagebox.showerror("Error", "Failed to save configuration")

    def skip_and_auto_detect():
        result[0] = "auto"
        root.destroy()

    save_btn = tk.Button(
        button_frame,
        text="Save & Continue",
        command=save_and_continue,
        width=15,
        font=("Arial", 10, "bold"),
        bg="#27ae60",
        fg="white",
    )
    save_btn.pack(side="left", padx=10)

    auto_btn = tk.Button(
        button_frame,
        text="Auto-Detect Files",
        command=skip_and_auto_detect,
        width=15,
        font=("Arial", 10),
        bg="#95a5a6",
        fg="white",
    )
    auto_btn.pack(side="left", padx=10)

    # Center window
    root.update_idletasks()
    x = (root.winfo_screenwidth() // 2) - (root.winfo_width() // 2)
    y = (root.winfo_screenheight() // 2) - (root.winfo_height() // 2)
    root.geometry(f"+{x}+{y}")

    root.mainloop()

    return result[0]


def auto_detect_config(workspace_path):
    """Automatically detect configuration from files in workspace"""
    import re
    from pathlib import Path

    workspace = Path(workspace_path)
    logger.debug(f"auto_detect_config: workspace_path={workspace_path}")

    # Check for input_files directory first, fallback to workspace root
    input_dir = workspace / "input_files"
    search_dir = input_dir if input_dir.exists() else workspace
    logger.debug(
        f"auto_detect_config: input_dir={input_dir}, exists={input_dir.exists()}"
    )
    logger.debug(f"auto_detect_config: search_dir={search_dir}")

    print(f"Searching for input files in: {search_dir}")  # Find BOM files
    bom_files = list(search_dir.glob("*Bill of Materials*.xlsx"))
    schematic_files = list(search_dir.glob("*Schematic*.PDF"))
    assembly_files = list(search_dir.glob("*Assembly*.PDF"))

    logger.debug(f"auto_detect_config: bom_files={[f.name for f in bom_files]}")
    logger.debug(
        f"auto_detect_config: schematic_files={[f.name for f in schematic_files]}"
    )
    logger.debug(
        f"auto_detect_config: assembly_files={[f.name for f in assembly_files]}"
    )

    print(
        f"Found {len(bom_files)} BOM file(s), {len(schematic_files)} schematic file(s), {len(assembly_files)} assembly file(s)"
    )

    if len(bom_files) >= 2:
        # Extract metadata from BOM file contents
        bom_metadata_list = []
        for bom_file in bom_files:
            metadata = extract_bom_metadata(bom_file)
            # Try to extract version from metadata, fallback to filename
            version_num = 0
            if metadata["version"] is not None:
                version_num = (
                    int(metadata["version"]) if metadata["version"].isdigit() else 0
                )
            else:
                # Fallback: Extract version from filename (e.g., "BT3413A-8")
                match = re.search(r"-(\d+)", bom_file.name)
                if match:
                    version_num = int(match.group(1))

            bom_metadata_list.append((bom_file, metadata, version_num))

        if len(bom_metadata_list) >= 2:
            # Sort by version number
            bom_metadata_list.sort(key=lambda x: x[2])

            bom_old, metadata_old, _ = bom_metadata_list[0]
            bom_new, metadata_new, _ = bom_metadata_list[1]

            CONFIG["bom_old_filename"] = bom_old.name
            CONFIG["bom_new_filename"] = bom_new.name

            # Use product name from BOM contents
            CONFIG["product_name"] = (
                metadata_old["product_name"] or metadata_old["assembly_no"][:7]
                if metadata_old["assembly_no"]
                else "Unknown"
            )

            # Use version from BOM contents or filename
            CONFIG["version_old"] = (
                f"V{metadata_old['version']}"
                if metadata_old["version"]
                else f"V{bom_metadata_list[0][2]}"
            )
            CONFIG["version_new"] = (
                f"V{metadata_new['version']}"
                if metadata_new["version"]
                else f"V{bom_metadata_list[1][2]}"
            )

            # Store variant information
            CONFIG["variant_old"] = metadata_old["variant"] or ""
            CONFIG["variant_letter_old"] = metadata_old["variant_letter"] or ""
            CONFIG["variant_new"] = metadata_new["variant"] or ""
            CONFIG["variant_letter_new"] = metadata_new["variant_letter"] or ""

    if len(schematic_files) >= 2:
        # Sort schematic files by version number
        pattern = r"-(\d+)"
        schematic_versions = []
        for schematic_file in schematic_files:
            match = re.search(pattern, schematic_file.name)
            if match:
                version_num = int(match.group(1))
                schematic_versions.append((schematic_file, version_num))

        if len(schematic_versions) >= 2:
            schematic_versions.sort(key=lambda x: x[1])
            CONFIG["schematic_old_filename"] = schematic_versions[0][0].name
            CONFIG["schematic_new_filename"] = schematic_versions[1][0].name

    if len(assembly_files) >= 2:
        # Sort assembly files by version number
        pattern = r"-(\d+)"
        assembly_versions = []
        for assembly_file in assembly_files:
            match = re.search(pattern, assembly_file.name)
            if match:
                version_num = int(match.group(1))
                assembly_versions.append((assembly_file, version_num))

        if len(assembly_versions) >= 2:
            assembly_versions.sort(key=lambda x: x[1])
            CONFIG["assembly_old_filename"] = assembly_versions[0][0].name
            CONFIG["assembly_new_filename"] = assembly_versions[1][0].name


# ============================================================================


def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    new_run.append(rPr)
    new_run.text = text

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)  # type: ignore


def create_header_footer(doc, company_name="LHA Systems Pty Ltd"):
    """Create professional header and footer for all pages"""
    # Access the first section
    section = doc.sections[0]

    # Header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Company name in header
    header_run = header_para.add_run(company_name)
    header_run.font.bold = True
    header_run.font.size = Pt(12)
    header_run.font.color.rgb = RGBColor(0, 0, 0)

    # Add horizontal line below header
    header_border = header_para._element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    header_border.append(borders)

    # Footer
    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=3, width=Inches(6.5))
    footer_table.autofit = False

    # Left cell - Document title
    left_cell = footer_table.rows[0].cells[0]
    left_para = left_cell.paragraphs[0]
    left_para.text = "Certificate of Conformity"
    left_para.runs[0].font.size = Pt(9)

    # Center cell - Page number
    center_cell = footer_table.rows[0].cells[1]
    center_para = center_cell.paragraphs[0]
    center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center_run = center_para.add_run("Page ")

    # Add page number field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")

    # Add text element for field result
    t1 = OxmlElement("w:t")
    t1.text = "1"

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")

    center_run._r.append(fldChar1)
    center_run._r.append(instrText)
    center_run._r.append(fldChar2)
    center_run._r.append(t1)
    center_run._r.append(fldChar3)

    center_para.add_run(" of ")

    # Add total pages field
    total_run = center_para.add_run()
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "begin")

    instrText2 = OxmlElement("w:instrText")
    instrText2.set(qn("xml:space"), "preserve")
    instrText2.text = "NUMPAGES"

    fldChar5 = OxmlElement("w:fldChar")
    fldChar5.set(qn("w:fldCharType"), "separate")

    # Add text element for field result
    t2 = OxmlElement("w:t")
    t2.text = "1"

    fldChar6 = OxmlElement("w:fldChar")
    fldChar6.set(qn("w:fldCharType"), "end")

    total_run._r.append(fldChar4)
    total_run._r.append(instrText2)
    total_run._r.append(fldChar5)
    total_run._r.append(t2)
    total_run._r.append(fldChar6)

    center_para.runs[0].font.size = Pt(9)

    # Right cell - Date and time
    right_cell = footer_table.rows[0].cells[2]
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para.text = datetime.now().strftime("%Y-%m-%d %H:%M")
    right_para.runs[0].font.size = Pt(9)


class COCReportGenerator:
    def __init__(self, workspace_path):
        self.workspace_path = Path(workspace_path)
        self.changes = {
            "bom_added": [],
            "bom_removed": [],
            "bom_modified": [],
            "schematic_notes": [],
            "assembly_notes": [],
            "assembly_v8_analysis": {},
            "assembly_v10_analysis": {},
            "smd_migration": [],
        }
        self.pdf_paths = {}
        self.change_reasons = {}
        self.reasons_file = self.workspace_path / "component_change_reasons.json"
        self.document_info = {}  # Store creator and timestamps
        self.load_change_reasons()

    def load_change_reasons(self):
        """Load previously saved change reasons from JSON file"""
        if self.reasons_file.exists():
            try:
                with open(self.reasons_file, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    # Convert ADDITIONAL_NOTES from array to string if needed
                    if "ADDITIONAL_NOTES" in data and isinstance(
                        data["ADDITIONAL_NOTES"], list
                    ):
                        data["ADDITIONAL_NOTES"] = "\n".join(data["ADDITIONAL_NOTES"])
                    self.change_reasons = data
                print(f"✓ Loaded {len(self.change_reasons)} saved change reasons")
            except Exception as e:  # noqa: BLE001
                print(f"⚠ Could not load change reasons: {e}")
                self.change_reasons = {}
        else:
            self.change_reasons = {}

    def save_change_reasons(self):
        """Save change reasons to JSON file with readable formatting"""
        try:
            # Manually format the JSON for better readability
            with open(self.reasons_file, "w", encoding="utf-8") as f:
                f.write("{\n")
                items = list(self.change_reasons.items())

                for idx, (key, value) in enumerate(items):
                    is_last = idx == len(items) - 1

                    if key == "ADDITIONAL_NOTES" and value:
                        # Format ADDITIONAL_NOTES across multiple lines
                        f.write(f'  "{key}": [\n')
                        lines = []
                        current_line = ""

                        for paragraph in value.split("\n"):
                            if not paragraph.strip():
                                if current_line:
                                    lines.append(current_line)
                                    current_line = ""
                                lines.append("")
                            else:
                                words = paragraph.split()
                                for word in words:
                                    if len(current_line) + len(word) + 1 <= 90:
                                        current_line += (
                                            " " if current_line else ""
                                        ) + word
                                    else:
                                        if current_line:
                                            lines.append(current_line)
                                        current_line = word

                        if current_line:
                            lines.append(current_line)

                        for line_idx, line in enumerate(lines):
                            is_last_line = line_idx == len(lines) - 1
                            f.write(f'    "{line}"')
                            if not is_last_line:
                                f.write(",\n")
                            else:
                                f.write("\n")

                        f.write("  ]")
                    else:
                        # Regular key-value pair
                        f.write(f'  "{key}": "{value}"')

                    if not is_last:
                        f.write(",\n")
                    else:
                        f.write("\n")

                f.write("}\n")

            print(f"✓ Saved {len(self.change_reasons)} change reasons")
        except Exception as e:  # noqa: BLE001
            print(f"⚠ Could not save change reasons: {e}")

    def prompt_for_change_reasons(self):
        """Interactive GUI questionnaire for component change reasons"""
        all_components = []

        # Collect all components that changed
        for item in self.changes["bom_added"]:
            key = f"ADDED_{item['part']}"
            all_components.append(("Added", item, key))

        for item in self.changes["bom_removed"]:
            key = f"REMOVED_{item['part']}"
            all_components.append(("Removed", item, key))

        for item in self.changes["bom_modified"]:
            key = f"MODIFIED_{item['part']}"
            all_components.append(("Modified", item, key))

        if not all_components:
            print("No component changes detected. Skipping questionnaire.")
            return

        # Create GUI window
        root = tk.Tk()
        root.title("Component Change Reasons")
        root.geometry("750x800")

        # Bring window to front
        root.lift()
        root.attributes("-topmost", True)
        root.after_idle(root.attributes, "-topmost", False)
        root.focus_force()

        # Track current component index and screen state
        current_index = [0]
        current_screen = ["creator_info"]  # Start with creator info

        # Header
        header_frame = tk.Frame(root, bg="#2c3e50", height=60)
        header_frame.pack(fill="x")
        header_frame.pack_propagate(False)

        header_label = tk.Label(
            header_frame,
            text="Component Change Questionnaire",
            font=("Arial", 16, "bold"),
            bg="#2c3e50",
            fg="white",
        )
        header_label.pack(pady=15)

        # Create canvas and scrollbar for scrollable content
        canvas = tk.Canvas(root)
        scrollbar = tk.Scrollbar(root, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True, padx=10, pady=10)
        scrollbar.pack(side="right", fill="y")

        # Main content frame inside scrollable area
        content_frame = tk.Frame(scrollable_frame, padx=20, pady=20)
        content_frame.pack(fill="both", expand=True)

        def show_creator_info():
            """Show document creator information and signature screen"""
            for widget in content_frame.winfo_children():
                widget.destroy()

            # Version Information Frame
            version_frame = tk.LabelFrame(
                content_frame,
                text="Product Version Information",
                font=("Arial", 11, "bold"),
                padx=20,
                pady=15,
                bg="#ecf0f1",
            )
            version_frame.grid(row=0, column=0, sticky="ew", pady=(0, 10))
            version_frame.columnconfigure((0, 1, 2), weight=1)

            # Product Name
            product_label = tk.Label(
                version_frame,
                text=f"Product: {CONFIG['product_name']}",
                font=("Arial", 12, "bold"),
                bg="#ecf0f1",
            )
            product_label.grid(row=0, column=0, columnspan=3, pady=(0, 10))

            # Old Version Column
            old_header = tk.Label(
                version_frame,
                text="OLD VERSION",
                font=("Arial", 10, "bold"),
                bg="#e74c3c",
                fg="white",
                padx=10,
                pady=5,
            )
            old_header.grid(row=1, column=0, sticky="ew", padx=(0, 5))

            old_version = tk.Label(
                version_frame,
                text=f"Version: {CONFIG['version_old']}",
                font=("Arial", 9),
                bg="#ecf0f1",
                anchor="w",
            )
            old_version.grid(row=2, column=0, sticky="ew", padx=(0, 5), pady=2)

            if CONFIG.get("variant_old"):
                old_variant = tk.Label(
                    version_frame,
                    text=f"Variant: {CONFIG['variant_old']} {CONFIG.get('variant_letter_old', '')}",
                    font=("Arial", 9),
                    bg="#ecf0f1",
                    anchor="w",
                )
                old_variant.grid(row=3, column=0, sticky="ew", padx=(0, 5), pady=2)

            old_bom = tk.Label(
                version_frame,
                text=f"BOM: {CONFIG['bom_old_filename'][:30]}...",
                font=("Arial", 8),
                bg="#ecf0f1",
                fg="#555",
                anchor="w",
            )
            old_bom.grid(row=4, column=0, sticky="ew", padx=(0, 5), pady=2)

            # Arrow/Change Indicator
            arrow_label = tk.Label(
                version_frame,
                text="→",
                font=("Arial", 24, "bold"),
                bg="#ecf0f1",
                fg="#3498db",
            )
            arrow_label.grid(row=1, column=1, rowspan=4, padx=5)

            # New Version Column
            new_header = tk.Label(
                version_frame,
                text="NEW VERSION",
                font=("Arial", 10, "bold"),
                bg="#27ae60",
                fg="white",
                padx=10,
                pady=5,
            )
            new_header.grid(row=1, column=2, sticky="ew", padx=(5, 0))

            new_version = tk.Label(
                version_frame,
                text=f"Version: {CONFIG['version_new']}",
                font=("Arial", 9),
                bg="#ecf0f1",
                anchor="w",
            )
            new_version.grid(row=2, column=2, sticky="ew", padx=(5, 0), pady=2)

            if CONFIG.get("variant_new"):
                new_variant = tk.Label(
                    version_frame,
                    text=f"Variant: {CONFIG['variant_new']} {CONFIG.get('variant_letter_new', '')}",
                    font=("Arial", 9),
                    bg="#ecf0f1",
                    anchor="w",
                )
                new_variant.grid(row=3, column=2, sticky="ew", padx=(5, 0), pady=2)

            new_bom = tk.Label(
                version_frame,
                text=f"BOM: {CONFIG['bom_new_filename'][:30]}...",
                font=("Arial", 8),
                bg="#ecf0f1",
                fg="#555",
                anchor="w",
            )
            new_bom.grid(row=4, column=2, sticky="ew", padx=(5, 0), pady=2)

            # Single grouped frame for all information
            info_frame = tk.LabelFrame(
                content_frame,
                text="Document Information & Approval Signatures",
                font=("Arial", 11, "bold"),
                padx=20,
                pady=15,
            )
            info_frame.grid(row=1, column=0, sticky="ew", pady=10)
            info_frame.columnconfigure(1, weight=1)

            # Created by field
            creator_label = tk.Label(
                info_frame, text="Created by:", font=("Arial", 10, "bold")
            )
            creator_label.grid(row=0, column=0, sticky="w", pady=8)

            creator_entry = tk.Entry(info_frame, width=50, font=("Arial", 10))
            creator_entry.grid(row=0, column=1, sticky="ew", pady=8, padx=(10, 0))

            # Load existing value if available
            existing_creator = self.change_reasons.get("DOCUMENT_CREATOR", "")
            creator_entry.insert(0, existing_creator)

            # Date/time display (auto-generated)
            datetime_label = tk.Label(
                info_frame, text="Start Date/Time:", font=("Arial", 10, "bold")
            )
            datetime_label.grid(row=1, column=0, sticky="w", pady=8)

            start_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            datetime_value = tk.Label(
                info_frame, text=start_datetime, font=("Arial", 10)
            )
            datetime_value.grid(row=1, column=1, sticky="w", pady=8, padx=(10, 0))

            # Separator
            separator = tk.Frame(info_frame, height=2, bg="#bdc3c7")
            separator.grid(row=2, column=0, columnspan=2, sticky="ew", pady=15)

            # Prepared by
            prepared_label = tk.Label(
                info_frame, text="Prepared by:", font=("Arial", 10)
            )
            prepared_label.grid(row=3, column=0, sticky="w", pady=8)

            prepared_entry = tk.Entry(info_frame, width=50, font=("Arial", 10))
            prepared_entry.grid(row=3, column=1, sticky="ew", pady=8, padx=(10, 0))
            prepared_entry.insert(0, self.change_reasons.get("PREPARED_BY", ""))

            # Reviewed by
            reviewed_label = tk.Label(
                info_frame, text="Reviewed by:", font=("Arial", 10)
            )
            reviewed_label.grid(row=4, column=0, sticky="w", pady=8)

            reviewed_entry = tk.Entry(info_frame, width=50, font=("Arial", 10))
            reviewed_entry.grid(row=4, column=1, sticky="ew", pady=8, padx=(10, 0))
            reviewed_entry.insert(0, self.change_reasons.get("REVIEWED_BY", ""))

            # Approved by
            approved_label = tk.Label(
                info_frame, text="Approved by:", font=("Arial", 10)
            )
            approved_label.grid(row=5, column=0, sticky="w", pady=8)

            approved_entry = tk.Entry(info_frame, width=50, font=("Arial", 10))
            approved_entry.grid(row=5, column=1, sticky="ew", pady=8, padx=(10, 0))
            approved_entry.insert(0, self.change_reasons.get("APPROVED_BY", ""))

            # Separator
            separator2 = tk.Frame(info_frame, height=2, bg="#bdc3c7")
            separator2.grid(row=6, column=0, columnspan=2, sticky="ew", pady=15)

            # Place/Address
            place_label = tk.Label(
                info_frame, text="Place/Address:", font=("Arial", 10, "bold")
            )
            place_label.grid(row=7, column=0, sticky="w", pady=8)

            place_entry = tk.Entry(info_frame, width=50, font=("Arial", 10))
            place_entry.grid(row=7, column=1, sticky="ew", pady=8, padx=(10, 0))
            place_entry.insert(0, self.change_reasons.get("DOCUMENT_PLACE", ""))

            def proceed_to_components():
                creator_name = creator_entry.get().strip()
                if not creator_name:
                    tk.messagebox.showwarning(
                        "Required Field",
                        "Please enter the creator name before proceeding.",
                    )
                    return

                self.change_reasons["DOCUMENT_CREATOR"] = creator_name
                self.change_reasons["DOCUMENT_START_TIMESTAMP"] = start_datetime

                # Save signatures
                self.change_reasons["PREPARED_BY"] = prepared_entry.get().strip()
                self.change_reasons["REVIEWED_BY"] = reviewed_entry.get().strip()
                self.change_reasons["APPROVED_BY"] = approved_entry.get().strip()
                self.change_reasons["DOCUMENT_PLACE"] = place_entry.get().strip()

                current_screen[0] = "components"
                show_component(0)

            # Button frame
            btn_frame = tk.Frame(content_frame)
            btn_frame.grid(row=2, column=0, pady=30)

            next_btn = tk.Button(
                btn_frame,
                text="Next",
                command=proceed_to_components,
                width=12,
                font=("Arial", 10),
                bg="#3498db",
                fg="white",
            )
            next_btn.pack()

            creator_entry.focus_set()

        def show_component(index):
            """Display the component at the given index"""
            if index < 0 or index >= len(all_components):
                return

            # Clear content frame
            for widget in content_frame.winfo_children():
                widget.destroy()

            change_type, item, key = all_components[index]
            part = item["part"]
            desc = item.get("description", "N/A")
            desig = item.get("designator", "N/A")

            # Create component info labels
            title_label = tk.Label(
                content_frame,
                text=f"{change_type} Component ({index + 1} of {len(all_components)})",
                font=("Arial", 12, "bold"),
            )
            title_label.grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 15))

            part_label = tk.Label(
                content_frame, text=f"Part Number: {part}", font=("Arial", 10)
            )
            part_label.grid(row=1, column=0, columnspan=2, sticky="w", padx=20, pady=5)

            desig_label = tk.Label(
                content_frame, text=f"Designator: {desig}", font=("Arial", 10)
            )
            desig_label.grid(row=2, column=0, columnspan=2, sticky="w", padx=20, pady=5)

            desc_label = tk.Label(
                content_frame,
                text=f"Description: {desc}",
                font=("Arial", 10),
                wraplength=600,
            )
            desc_label.grid(row=3, column=0, columnspan=2, sticky="w", padx=20, pady=5)

            if change_type == "Modified":
                qty_label = tk.Label(
                    content_frame,
                    text=f"Quantity Change: {item['old_qty']} → {item['new_qty']}",
                    font=("Arial", 10),
                )
                qty_label.grid(
                    row=4, column=0, columnspan=2, sticky="w", padx=20, pady=5
                )

            # Reason input
            reason_label = tk.Label(
                content_frame, text="Reason for change:", font=("Arial", 10, "bold")
            )
            reason_label.grid(row=5, column=0, columnspan=2, sticky="w", pady=(20, 5))

            reason_text = scrolledtext.ScrolledText(
                content_frame, height=8, width=70, font=("Arial", 10)
            )
            reason_text.grid(row=6, column=0, columnspan=2, sticky="ew", pady=(0, 20))

            # Load previous reason if exists
            previous_reason = self.change_reasons.get(key, "")
            reason_text.insert("1.0", previous_reason)

            def save_current_reason():
                """Save the reason for the current component"""
                reason = reason_text.get("1.0", tk.END).strip()
                if reason:
                    self.change_reasons[key] = reason
                elif key not in self.change_reasons:
                    self.change_reasons[key] = "Not specified"

            def prev_component():
                """Go to previous component"""
                save_current_reason()
                current_index[0] -= 1
                show_component(current_index[0])

            def next_component():
                """Go to next component or show additional notes"""
                save_current_reason()
                if current_index[0] < len(all_components) - 1:
                    current_index[0] += 1
                    show_component(current_index[0])
                else:
                    # Show additional notes screen
                    show_additional_notes()

            # Button frame
            button_frame = tk.Frame(content_frame)
            button_frame.grid(row=7, column=0, columnspan=2, pady=10)

            prev_btn = tk.Button(
                button_frame,
                text="Previous",
                command=prev_component,
                width=12,
                font=("Arial", 10),
                state=tk.NORMAL if index > 0 else tk.DISABLED,
            )
            prev_btn.pack(side="left", padx=10)

            next_btn = tk.Button(
                button_frame,
                text="Next" if index < len(all_components) - 1 else "Finish",
                command=next_component,
                width=12,
                font=("Arial", 10),
                bg="#3498db",
                fg="white",
            )
            next_btn.pack(side="left", padx=10)

            skip_btn = tk.Button(
                button_frame,
                text="Skip All",
                command=lambda: [self.save_change_reasons(), root.destroy()],
                width=12,
                font=("Arial", 10),
            )
            skip_btn.pack(side="left", padx=10)

        def show_additional_notes():
            """Show additional notes screen after all components"""
            # Clear content frame
            for widget in content_frame.winfo_children():
                widget.destroy()

            # Show engineering summary screen
            notes_title = tk.Label(
                content_frame, text="Engineering Summary", font=("Arial", 12, "bold")
            )
            notes_title.grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 15))

            notes_instruction = tk.Label(
                content_frame,
                text="Add engineering summary or overall comments about the changes:",
                font=("Arial", 10),
            )
            notes_instruction.grid(
                row=1, column=0, columnspan=2, sticky="w", pady=(0, 10)
            )

            notes_text_widget = scrolledtext.ScrolledText(
                content_frame, height=10, width=70, font=("Arial", 10)
            )
            notes_text_widget.grid(
                row=2, column=0, columnspan=2, sticky="ew", pady=(0, 20)
            )

            # Load existing notes if any
            existing_notes = self.change_reasons.get("ADDITIONAL_NOTES", "")
            notes_text_widget.insert("1.0", existing_notes)

            def finish_questionnaire():
                """Save additional notes and close"""
                notes = notes_text_widget.get("1.0", tk.END).strip()
                if notes:
                    self.change_reasons["ADDITIONAL_NOTES"] = notes

                self.save_change_reasons()
                root.destroy()

            def go_back():
                """Go back to last component"""
                current_index[0] = len(all_components) - 1
                show_component(current_index[0])

            # Finish button
            finish_button_frame = tk.Frame(content_frame)
            finish_button_frame.grid(row=3, column=0, columnspan=2, pady=10)

            back_btn = tk.Button(
                finish_button_frame,
                text="Back",
                command=go_back,
                width=12,
                font=("Arial", 10),
            )
            back_btn.pack(side="left", padx=10)

            finish_btn = tk.Button(
                finish_button_frame,
                text="Finish",
                command=finish_questionnaire,
                width=12,
                font=("Arial", 10),
                bg="#27ae60",
                fg="white",
            )
            finish_btn.pack(side="left", padx=10)

        # Show creator info screen first
        show_creator_info()

        # Enable mouse wheel scrolling
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        canvas.bind_all("<MouseWheel>", on_mousewheel)

        # Center window on screen
        root.update_idletasks()
        x = (root.winfo_screenwidth() // 2) - (root.winfo_width() // 2)
        y = (root.winfo_screenheight() // 2) - (root.winfo_height() // 2)
        root.geometry(f"+{x}+{y}")

        # Run the GUI
        root.mainloop()

        print(f"✓ Collected reasons for {len(self.change_reasons)} components")

    def load_bom(self, file_path):
        """Load BOM from Excel file"""
        try:
            # First, find the header row by looking for "SeqNum" or similar markers
            import openpyxl

            wb = openpyxl.load_workbook(file_path)
            ws = wb.active

            header_row_idx = None
            for idx, row in enumerate(ws.iter_rows(max_row=20, values_only=True)):
                if row and any(cell and "SeqNum" in str(cell) for cell in row):
                    header_row_idx = idx
                    break
            wb.close()

            # Default to row 10 if not found
            if header_row_idx is None:
                header_row_idx = 9  # 0-indexed, so row 10

            # Load with detected header row
            df = pd.read_excel(file_path, header=header_row_idx)
            # Drop rows with all NaN
            df = df.dropna(how="all")
            # Drop columns with all NaN
            df = df.dropna(axis=1, how="all")
            print(f"✓ Loaded BOM: {file_path.name}")
            print(f"  Columns: {list(df.columns)}")
            print(f"  Rows: {len(df)}")
            return df
        except Exception as e:  # noqa: BLE001
            print(f"✗ Error loading {file_path.name}: {e}")
            return None

    def compare_boms(self, bom_v8_path, bom_v10_path):
        """Compare two BOM versions"""
        print("\n" + "=" * 60)
        print("COMPARING BOMs")
        print("=" * 60)

        bom_v8 = self.load_bom(bom_v8_path)
        bom_v10 = self.load_bom(bom_v10_path)

        if bom_v8 is None or bom_v10 is None:
            return

        # Try to identify key columns
        possible_part_cols = [
            "Part Number",
            "Part No",
            "Part",
            "Item",
            "Component",
            "Reference",
        ]
        possible_desc_cols = ["Description", "Desc", "Name", "Component Name"]
        possible_qty_cols = ["Quantity", "Qty", "Q'ty", "Amount"]
        possible_desig_cols = [
            "Designator",
            "Designators",
            "Reference",
            "Ref",
            "RefDes",
        ]

        part_col = next(
            (
                col
                for col in bom_v8.columns
                if any(pc.lower() in col.lower() for pc in possible_part_cols)
            ),
            None,
        )
        desc_col = next(
            (
                col
                for col in bom_v8.columns
                if any(dc.lower() in col.lower() for dc in possible_desc_cols)
            ),
            None,
        )
        qty_col = next(
            (
                col
                for col in bom_v8.columns
                if any(qc.lower() in col.lower() for qc in possible_qty_cols)
            ),
            None,
        )
        desig_col = next(
            (
                col
                for col in bom_v8.columns
                if any(dc.lower() in col.lower() for dc in possible_desig_cols)
            ),
            None,
        )

        if not part_col:
            part_col = bom_v8.columns[0] if len(bom_v8.columns) > 0 else None

        print("\nIdentified columns:")
        print(f"  Part Number: {part_col}")
        print(f"  Description: {desc_col}")
        print(f"  Quantity: {qty_col}")
        print(f"  Designator: {desig_col}")

        # Create comparison keys
        if part_col:
            v8_parts = set(bom_v8[part_col].dropna().astype(str))
            v10_parts = set(bom_v10[part_col].dropna().astype(str))

            added = v10_parts - v8_parts
            removed = v8_parts - v10_parts
            common = v8_parts & v10_parts

            print("\n📊 BOM Summary:")
            print(f"  V8 parts: {len(v8_parts)}")
            print(f"  V10 parts: {len(v10_parts)}")
            print(f"  Added: {len(added)}")
            print(f"  Removed: {len(removed)}")
            print(f"  Common: {len(common)}")

            # Store added parts
            for part in added:
                row = bom_v10[bom_v10[part_col].astype(str) == part].iloc[0]
                desc = row[desc_col] if desc_col and desc_col in row else "N/A"
                qty = row[qty_col] if qty_col and qty_col in row else "N/A"
                desig = row[desig_col] if desig_col and desig_col in row else "N/A"
                # Format quantity without decimals if it's a whole number
                if qty != "N/A":
                    try:
                        qty_float = float(qty)
                        # Skip if quantity is 0
                        if qty_float == 0:
                            continue
                        qty = (
                            str(int(qty_float))
                            if qty_float.is_integer()
                            else str(qty_float)
                        )
                    except (ValueError, TypeError):
                        qty = str(qty)
                self.changes["bom_added"].append(
                    {
                        "part": part,
                        "description": desc,
                        "quantity": qty,
                        "designator": desig,
                    }
                )

            # Sort added components by designator
            self.changes["bom_added"] = sorted(
                self.changes["bom_added"], key=lambda x: str(x.get("designator", "ZZZ"))
            )

            # Store removed parts
            for part in removed:
                row = bom_v8[bom_v8[part_col].astype(str) == part].iloc[0]
                desc = row[desc_col] if desc_col and desc_col in row else "N/A"
                qty = row[qty_col] if qty_col and qty_col in row else "N/A"
                desig = row[desig_col] if desig_col and desig_col in row else "N/A"
                # Format quantity without decimals if it's a whole number
                if qty != "N/A":
                    try:
                        qty_float = float(qty)
                        # Skip if quantity is 0
                        if qty_float == 0:
                            continue
                        qty = (
                            str(int(qty_float))
                            if qty_float.is_integer()
                            else str(qty_float)
                        )
                    except (ValueError, TypeError):
                        qty = str(qty)
                self.changes["bom_removed"].append(
                    {
                        "part": part,
                        "description": desc,
                        "quantity": qty,
                        "designator": desig,
                    }
                )

            # Sort removed components by designator
            self.changes["bom_removed"] = sorted(
                self.changes["bom_removed"],
                key=lambda x: str(x.get("designator", "ZZZ")),
            )

            # Check for quantity changes in common parts
            if qty_col:
                for part in common:
                    v8_row = bom_v8[bom_v8[part_col].astype(str) == part].iloc[0]
                    v10_row = bom_v10[bom_v10[part_col].astype(str) == part].iloc[0]

                    try:
                        v8_qty = float(v8_row[qty_col])
                        v10_qty = float(v10_row[qty_col])

                        if v8_qty != v10_qty:
                            desc = (
                                v10_row[desc_col]
                                if desc_col and desc_col in v10_row
                                else "N/A"
                            )
                            desig = (
                                v10_row[desig_col]
                                if desig_col and desig_col in v10_row
                                else "N/A"
                            )
                            # Format without decimals if whole numbers
                            v8_qty_str = (
                                str(int(v8_qty)) if v8_qty.is_integer() else str(v8_qty)
                            )
                            v10_qty_str = (
                                str(int(v10_qty))
                                if v10_qty.is_integer()
                                else str(v10_qty)
                            )
                            self.changes["bom_modified"].append(
                                {
                                    "part": part,
                                    "description": desc,
                                    "designator": desig,
                                    "old_qty": v8_qty_str,
                                    "new_qty": v10_qty_str,
                                }
                            )
                    except (ValueError, TypeError, AttributeError):
                        pass

            # Sort modified components by designator
            self.changes["bom_modified"] = sorted(
                self.changes["bom_modified"],
                key=lambda x: str(x.get("designator", "ZZZ")),
            )

    def analyze_pdf(self, pdf_path, doc_type):
        """Extract text from PDF for analysis"""
        print(f"\n📄 Analyzing {doc_type}: {pdf_path.name}")
        try:
            with open(pdf_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()

                print(f"  ✓ Extracted {len(text)} characters")
                print(f"  Pages: {len(reader.pages)}")
                return text
        except Exception as e:  # noqa: BLE001
            print(f"  ✗ Error: {e}")
            return None

    def analyze_assembly_drawing(self, pdf_path, version):
        """Analyze assembly drawing for component placement"""
        print(f"\n🔍 Analyzing Assembly Drawing {version}: {pdf_path.name}")
        analysis = {
            "total_components": 0,
            "top_side": 0,
            "bottom_side": 0,
            "smd_components": 0,
            "through_hole": 0,
            "pages": 0,
        }

        try:
            # Extract text and analyze
            doc = fitz.open(pdf_path)
            analysis["pages"] = len(doc)

            all_text = ""
            for page in doc:
                page_text = page.get_text()
                if isinstance(page_text, str):
                    all_text += page_text

            # Look for component placement indicators
            top_keywords = ["TOP", "TOP SIDE", "COMPONENT SIDE", "PRIMARY SIDE"]
            bottom_keywords = ["BOTTOM", "BOTTOM SIDE", "SOLDER SIDE", "SECONDARY SIDE"]
            smd_keywords = [
                "SMD",
                "SMT",
                "SURFACE MOUNT",
                "0805",
                "0603",
                "1206",
                "SOT",
                "SOIC",
                "QFN",
            ]
            thl_keywords = ["THL", "THROUGH HOLE", "THT", "RADIAL", "AXIAL"]

            # Count references to sides
            for keyword in top_keywords:
                analysis["top_side"] += all_text.upper().count(keyword)

            for keyword in bottom_keywords:
                analysis["bottom_side"] += all_text.upper().count(keyword)

            for keyword in smd_keywords:
                analysis["smd_components"] += all_text.upper().count(keyword)

            for keyword in thl_keywords:
                analysis["through_hole"] += all_text.upper().count(keyword)

            # Extract component designators (C1, R1, U1, etc.)
            designators = re.findall(r"\b[CRUDQLK]\d+", all_text.upper())
            analysis["total_components"] = len(set(designators))

            print(f"  ✓ Components found: {analysis['total_components']}")
            print(f"  ✓ Top side references: {analysis['top_side']}")
            print(f"  ✓ Bottom side references: {analysis['bottom_side']}")
            print(f"  ✓ SMD references: {analysis['smd_components']}")
            print(f"  ✓ Through-hole references: {analysis['through_hole']}")

            doc.close()
            return analysis

        except Exception as e:  # noqa: BLE001
            print(f"  ✗ Error: {e}")
            return analysis

    def pdf_page_to_image(self, pdf_path, page_num=0):
        """Convert PDF page to image and return image path"""
        try:
            doc = fitz.open(pdf_path)
            page = doc[page_num]

            # Render page to pixmap (image)
            zoom = 2  # Higher zoom for better quality
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)

            # Convert to PIL Image
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))

            # Save temporarily
            temp_path = pdf_path.parent / f"temp_{pdf_path.stem}_p{page_num}.png"
            img.save(temp_path, "PNG")

            doc.close()
            return temp_path
        except Exception as e:  # noqa: BLE001
            print(f"  ✗ Error converting PDF to image: {e}")
            return None

    def compare_assembly_drawings(self):
        """Compare assembly drawings to identify SMD migration"""
        print("\n" + "=" * 60)
        print("ANALYZING ASSEMBLY DRAWINGS")
        print("=" * 60)

        v8 = self.changes["assembly_v8_analysis"]
        v10 = self.changes["assembly_v10_analysis"]

        # Detect SMD migration to top side
        if v10.get("top_side", 0) > v8.get("top_side", 0):
            self.changes["smd_migration"].append(
                {
                    "change": "SMD components migrated to top side",
                    "v8_top_refs": v8.get("top_side", 0),
                    "v10_top_refs": v10.get("top_side", 0),
                    "benefit": "Production speed up and efficiency",
                }
            )
            print("\n✓ Detected: SMD migration to top side")

        if v10.get("smd_components", 0) > v8.get("smd_components", 0):
            self.changes["smd_migration"].append(
                {
                    "change": "Increased SMD component usage",
                    "v8_smd": v8.get("smd_components", 0),
                    "v10_smd": v10.get("smd_components", 0),
                    "benefit": "Improved manufacturability",
                }
            )

    def generate_word_report(self, output_path):
        """Generate COC report in Word format"""
        print("\n" + "=" * 60)
        print("GENERATING COC REPORT")
        print("=" * 60)

        doc = Document()

        # Create header and footer
        create_header_footer(doc)

        # Company Logo
        logo_path = self.workspace_path / "LHA_logo.png"
        if logo_path.exists():
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_para.add_run()
            logo_run.add_picture(str(logo_path), width=Inches(2.5))

        # Company name
        company = doc.add_paragraph()
        company.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = company.add_run("LHA Systems (Pty) Ltd")
        company_run.font.size = Pt(16)
        company_run.font.bold = True

        # Tagline
        tagline = doc.add_paragraph()
        tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tagline_run = tagline.add_run("Electronic Design & Development")
        tagline_run.font.size = Pt(10)
        tagline_run.font.italic = True

        # Address Information
        doc.add_paragraph()
        address_heading = doc.add_paragraph()
        address_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address_heading_run = address_heading.add_run("Company Information")
        address_heading_run.font.size = Pt(12)
        address_heading_run.font.bold = True

        # Full company details
        address_para = doc.add_paragraph()
        address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address_lines = [
            "1 Innovation Centre II · Meson Street · Technopark",
            "STELLENBOSCH · 7600 · South Africa",
            "",
            "Phone: +27 21 880-1886 · Fax: +27 21 880-1211",
            "",
            "Directors: CW Bosch; CH Malan; LHA Rossouw (Managing)",
        ]
        for line in address_lines:
            address_para.add_run(line + "\n")

        doc.add_paragraph()

        # Title
        title = doc.add_heading("Certificate of Conformity (COC)", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Document info
        doc.add_heading("Document Information", 1)

        # Determine number of rows based on available data
        num_rows = 4  # Base rows: Product, Version, Report Date, Total Pages
        if "DOCUMENT_CREATOR" in self.change_reasons:
            num_rows += 1
        if "DOCUMENT_START_TIMESTAMP" in self.change_reasons:
            num_rows += 1

        info_table = doc.add_table(rows=num_rows, cols=2)
        info_table.style = "Light Grid Accent 1"

        info_data = [
            ["Product", CONFIG["product_name"]],
            [
                "Version Comparison",
                f"{CONFIG['version_old']} → {CONFIG['version_new']}",
            ],
            ["Report Date", datetime.now().strftime("%Y-%m-%d %H:%M")],
        ]

        # Add creator and start timestamp if available
        if "DOCUMENT_CREATOR" in self.change_reasons:
            info_data.append(
                ["Document Creator", self.change_reasons["DOCUMENT_CREATOR"]]
            )
        if "DOCUMENT_START_TIMESTAMP" in self.change_reasons:
            info_data.append(
                ["Started", self.change_reasons["DOCUMENT_START_TIMESTAMP"]]
            )

        for i, (key, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = key
            info_table.rows[i].cells[1].text = value

        # Add Total Pages row with field code
        pages_row_idx = len(info_data)
        info_table.rows[pages_row_idx].cells[0].text = "Total Pages"
        pages_cell = info_table.rows[pages_row_idx].cells[1]
        pages_para = pages_cell.paragraphs[0]

        # Add NUMPAGES field
        pages_run = pages_para.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "NUMPAGES"

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")

        # Add text element for field result
        t = OxmlElement("w:t")
        t.text = "1"

        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")

        pages_run._r.append(fldChar1)
        pages_run._r.append(instrText)
        pages_run._r.append(fldChar2)
        pages_run._r.append(t)
        pages_run._r.append(fldChar3)

        # Document Purpose
        doc.add_paragraph()
        doc.add_heading("Document Purpose", 1)

        purpose_para = doc.add_paragraph()
        purpose_text = (
            "This Certificate of Conformity (COC) documents the engineering changes "
            f"made between version {CONFIG['version_old']} and {CONFIG['version_new']} of the {CONFIG['product_name']} product. This document provides "
            "a comprehensive record of all Bill of Materials (BOM) modifications, including "
            "component additions, removals, and quantity changes. Each change is accompanied "
            "by technical justification and analysis of schematic and assembly drawing updates. "
            "This COC serves as an official record for quality assurance, regulatory compliance, "
            "and change management purposes, ensuring full traceability of product evolution "
            "and design decisions."
        )
        purpose_para.add_run(purpose_text)
        purpose_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # BOM Changes
        doc.add_heading("Bill of Materials Changes", 1)

        # Added Components
        if self.changes["bom_added"]:
            doc.add_heading("Added Components", 2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Set fixed column widths
            for row in table.rows:
                row.cells[0].width = Inches(1.2)
                row.cells[1].width = Inches(1.8)
                row.cells[2].width = Inches(3.0)
                row.cells[3].width = Inches(0.8)

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Designator"
            hdr_cells[1].text = "Part Number"
            hdr_cells[2].text = "Description"
            hdr_cells[3].text = "Quantity"
            hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for item in self.changes["bom_added"]:
                row = table.add_row()
                row.cells[0].width = Inches(1.2)
                row.cells[1].width = Inches(1.8)
                row.cells[2].width = Inches(3.0)
                row.cells[3].width = Inches(0.8)
                row.cells[0].text = str(item.get("designator", "N/A"))
                row.cells[1].text = str(item["part"])
                row.cells[2].text = str(item["description"])
                row.cells[3].text = str(item["quantity"])
                row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(
                f"Total added: {len(self.changes['bom_added'])} components"
            )

            # Add reasons for each component
            doc.add_paragraph()
            reasons_heading = doc.add_paragraph()
            reasons_heading.add_run("Reasons for Added Components:").bold = True

            for item in self.changes["bom_added"]:
                key = f"ADDED_{item['part']}"
                reason = self.change_reasons.get(key, "Not specified")
                reason_para = doc.add_paragraph(style="List Bullet")
                reason_para.add_run(
                    f"{item.get('designator', 'N/A')} ({item['part']}): "
                ).bold = True
                reason_para.add_run(reason)
        else:
            doc.add_paragraph("No components added.")

        # Removed Components
        if self.changes["bom_removed"]:
            doc.add_heading("Removed Components", 2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Set fixed column widths
            for row in table.rows:
                row.cells[0].width = Inches(1.2)
                row.cells[1].width = Inches(1.8)
                row.cells[2].width = Inches(3.0)
                row.cells[3].width = Inches(0.8)

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Designator"
            hdr_cells[1].text = "Part Number"
            hdr_cells[2].text = "Description"
            hdr_cells[3].text = "Quantity"
            hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for item in self.changes["bom_removed"]:
                row = table.add_row()
                row.cells[0].width = Inches(1.2)
                row.cells[1].width = Inches(1.8)
                row.cells[2].width = Inches(3.0)
                row.cells[3].width = Inches(0.8)
                row.cells[0].text = str(item.get("designator", "N/A"))
                row.cells[1].text = str(item["part"])
                row.cells[2].text = str(item["description"])
                row.cells[3].text = str(item["quantity"])
                row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(
                f"Total removed: {len(self.changes['bom_removed'])} components"
            )

            # Add reasons for each component
            doc.add_paragraph()
            reasons_heading = doc.add_paragraph()
            reasons_heading.add_run("Reasons for Removed Components:").bold = True

            for item in self.changes["bom_removed"]:
                key = f"REMOVED_{item['part']}"
                reason = self.change_reasons.get(key, "Not specified")
                reason_para = doc.add_paragraph(style="List Bullet")
                reason_para.add_run(
                    f"{item.get('designator', 'N/A')} ({item['part']}): "
                ).bold = True
                reason_para.add_run(reason)
        else:
            doc.add_paragraph("No components removed.")

        # Modified Components
        if self.changes["bom_modified"]:
            doc.add_heading("Modified Component Quantities", 2)
            table = doc.add_table(rows=1, cols=5)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Set fixed column widths
            for row in table.rows:
                row.cells[0].width = Inches(1.2)
                row.cells[1].width = Inches(1.8)
                row.cells[2].width = Inches(2.5)
                row.cells[3].width = Inches(0.8)
                row.cells[4].width = Inches(0.8)

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Designator"
            hdr_cells[1].text = "Part Number"
            hdr_cells[2].text = "Description"
            hdr_cells[3].text = "Old Qty (V8)"
            hdr_cells[4].text = "New Qty (V10)"
            hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for item in self.changes["bom_modified"]:
                row = table.add_row()
                row.cells[0].width = Inches(1.2)
                row.cells[1].width = Inches(1.8)
                row.cells[2].width = Inches(2.5)
                row.cells[3].width = Inches(0.8)
                row.cells[4].width = Inches(0.8)
                row.cells[0].text = str(item.get("designator", "N/A"))
                row.cells[1].text = str(item["part"])
                row.cells[2].text = str(item["description"])
                row.cells[3].text = str(item["old_qty"])
                row.cells[4].text = str(item["new_qty"])
                row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(
                f"Total modified: {len(self.changes['bom_modified'])} components"
            )

            # Add reasons for each component
            doc.add_paragraph()
            reasons_heading = doc.add_paragraph()
            reasons_heading.add_run("Reasons for Modified Components:").bold = True

            for item in self.changes["bom_modified"]:
                key = f"MODIFIED_{item['part']}"
                reason = self.change_reasons.get(key, "Not specified")
                reason_para = doc.add_paragraph(style="List Bullet")
                reason_para.add_run(
                    f"{item.get('designator', 'N/A')} ({item['part']}): "
                ).bold = True
                reason_para.add_run(reason)
        else:
            doc.add_paragraph("No quantity changes.")

        # Schematic Analysis
        doc.add_heading("Schematic Analysis", 1)
        doc.add_paragraph(
            "Schematic PDFs analyzed. Manual review recommended for circuit changes."
        )
        doc.add_paragraph("Files compared:")

        if "schematic_v8" in self.pdf_paths:
            p = doc.add_paragraph("  • ", style="List Bullet")
            add_hyperlink(
                p,
                str(self.pdf_paths["schematic_v8"]),
                self.pdf_paths["schematic_v8"].name,
            )
        if "schematic_v10" in self.pdf_paths:
            p = doc.add_paragraph("  • ", style="List Bullet")
            add_hyperlink(
                p,
                str(self.pdf_paths["schematic_v10"]),
                self.pdf_paths["schematic_v10"].name,
            )

        # Add schematic images
        doc.add_heading("Schematic Comparison (Page 1)", 2)

        if "schematic_v8" in self.pdf_paths and self.pdf_paths["schematic_v8"].exists():
            doc.add_paragraph("Version 8:", style="Heading 3")
            img_path = self.pdf_page_to_image(self.pdf_paths["schematic_v8"], 0)
            if img_path and img_path.exists():
                doc.add_picture(str(img_path), width=Inches(6.0))
                img_path.unlink()  # Delete temp file

        if (
            "schematic_v10" in self.pdf_paths
            and self.pdf_paths["schematic_v10"].exists()
        ):
            doc.add_paragraph("Version 10:", style="Heading 3")
            img_path = self.pdf_page_to_image(self.pdf_paths["schematic_v10"], 0)
            if img_path and img_path.exists():
                doc.add_picture(str(img_path), width=Inches(6.0))
                img_path.unlink()  # Delete temp file

        # Assembly Drawing Analysis
        doc.add_heading("Assembly Drawing Analysis", 1)

        # Add PDF links
        doc.add_paragraph("Files compared:")
        if "assembly_v8" in self.pdf_paths:
            p = doc.add_paragraph("  • ", style="List Bullet")
            add_hyperlink(
                p,
                str(self.pdf_paths["assembly_v8"]),
                self.pdf_paths["assembly_v8"].name,
            )
        if "assembly_v10" in self.pdf_paths:
            p = doc.add_paragraph("  • ", style="List Bullet")
            add_hyperlink(
                p,
                str(self.pdf_paths["assembly_v10"]),
                self.pdf_paths["assembly_v10"].name,
            )

        # Component placement analysis
        if (
            self.changes["assembly_v8_analysis"]
            or self.changes["assembly_v10_analysis"]
        ):
            doc.add_heading("Component Placement Analysis", 2)

            comp_table = doc.add_table(rows=1, cols=3)
            comp_table.style = "Light Grid Accent 1"
            hdr_cells = comp_table.rows[0].cells
            hdr_cells[0].text = "Metric"
            hdr_cells[1].text = "Version 8"
            hdr_cells[2].text = "Version 10"

            v8 = self.changes["assembly_v8_analysis"]
            v10 = self.changes["assembly_v10_analysis"]

            metrics = [
                ("Top Side References", v8.get("top_side", 0), v10.get("top_side", 0)),
                (
                    "Bottom Side References",
                    v8.get("bottom_side", 0),
                    v10.get("bottom_side", 0),
                ),
                (
                    "SMD References",
                    v8.get("smd_components", 0),
                    v10.get("smd_components", 0),
                ),
                (
                    "Through-Hole References",
                    v8.get("through_hole", 0),
                    v10.get("through_hole", 0),
                ),
            ]

            for metric, v8_val, v10_val in metrics:
                row_cells = comp_table.add_row().cells
                row_cells[0].text = metric
                row_cells[1].text = str(v8_val)
                row_cells[2].text = str(v10_val)

        # SMD Migration Notes
        if self.changes["smd_migration"]:
            doc.add_heading("Manufacturing Improvements", 2)
            doc.add_paragraph(
                "✓ SMD components migrated to top side for production speed up and efficiency",
                style="List Bullet",
            )
            for item in self.changes["smd_migration"]:
                p = doc.add_paragraph(
                    f"  - {item['change']}: {item['benefit']}", style="List Bullet 2"
                )
        else:
            doc.add_paragraph(
                "Assembly drawings analyzed. Manual review recommended for layout changes."
            )

        # Add assembly drawing images
        doc.add_heading("Assembly Drawing Comparison (Page 1)", 2)

        if "assembly_v8" in self.pdf_paths and self.pdf_paths["assembly_v8"].exists():
            doc.add_paragraph("Version 8:", style="Heading 3")
            img_path = self.pdf_page_to_image(self.pdf_paths["assembly_v8"], 0)
            if img_path and img_path.exists():
                doc.add_picture(str(img_path), width=Inches(6.0))
                img_path.unlink()  # Delete temp file

        if "assembly_v10" in self.pdf_paths and self.pdf_paths["assembly_v10"].exists():
            doc.add_paragraph("Version 10:", style="Heading 3")
            img_path = self.pdf_page_to_image(self.pdf_paths["assembly_v10"], 0)
            if img_path and img_path.exists():
                doc.add_picture(str(img_path), width=Inches(6.0))
                img_path.unlink()  # Delete temp file

        # Summary
        doc.add_heading("Summary", 1)
        summary = doc.add_paragraph()
        summary.add_run("Total Changes:\n").bold = True
        summary.add_run(f"  • Components Added: {len(self.changes['bom_added'])}\n")
        summary.add_run(f"  • Components Removed: {len(self.changes['bom_removed'])}\n")
        summary.add_run(
            f"  • Quantities Modified: {len(self.changes['bom_modified'])}\n"
        )

        # Engineering Summary
        if (
            "ADDITIONAL_NOTES" in self.change_reasons
            and self.change_reasons["ADDITIONAL_NOTES"].strip()
        ):
            doc.add_heading("Engineering Change Summary", 1)

            # Main summary paragraph
            summary_para = doc.add_paragraph()
            summary_text = (
                f"The modifications implemented in version {CONFIG['version_new']} represent a series of carefully planned "
                "engineering improvements focused on manufacturing optimization, electromagnetic compatibility, "
                "and cosmetic enhancements. All changes have been executed with strict adherence to maintaining "
                f"form, fit, and functional equivalence with the previous {CONFIG['version_old']} revision."
            )
            summary_para.add_run(summary_text)
            summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            doc.add_paragraph()

            # Key improvement areas heading
            improvements_heading = doc.add_paragraph()
            improvements_heading.add_run("Key Improvement Areas:").bold = True

            # Bullet points for key areas
            improvement_points = [
                (
                    "Electromagnetic Interference (EMI) and Emissions Optimization: ",
                    "Component selections and layout modifications have been implemented to enhance electromagnetic "
                    "compatibility, reduce radiated emissions, and improve immunity to external interference. These "
                    "changes ensure compliance with applicable EMC standards while maintaining identical electrical "
                    "performance characteristics.",
                ),
                (
                    "Manufacturing Process Enhancement: ",
                    "The PCB design has been refined to streamline assembly processes, improve component placement "
                    "accessibility, and reduce manufacturing complexity. These optimizations facilitate more efficient "
                    "production workflows, improved quality control, and reduced assembly time without affecting the "
                    "final product specifications or performance.",
                ),
                (
                    "Cosmetic and Board Layout Refinements: ",
                    "Visual and organizational improvements to the PCB layout enhance product aesthetics, improve "
                    "component identification, and optimize thermal management. These refinements contribute to a more "
                    "professional appearance and simplified maintenance procedures while preserving all functional "
                    "requirements.",
                ),
                (
                    "Form, Fit, and Function Equivalence: ",
                    f"It is hereby confirmed that version {CONFIG['version_new']} maintains complete form, fit, and functional equivalence "
                    f"with version {CONFIG['version_old']}. All mechanical dimensions, mounting interfaces, electrical connections, input/output "
                    "specifications, and operational characteristics remain unchanged. The product is fully interchangeable "
                    "with the previous version and requires no modifications to existing system integration, installation "
                    "procedures, or operational protocols.",
                ),
            ]

            for title, description in improvement_points:
                point_para = doc.add_paragraph(style="List Bullet")
                point_run = point_para.add_run(title)
                point_run.bold = True
                point_para.add_run(description)
                point_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Document Approval Signatures
        doc.add_heading("Document Approval", 1)

        # Add instruction
        doc.add_paragraph(
            "The undersigned hereby approve this Certificate of Conformity document."
        )

        # Prepared by
        doc.add_paragraph()
        prepared_para = doc.add_paragraph()
        prepared_para.add_run("Prepared by: ").bold = True
        prepared_para.add_run(self.change_reasons.get("PREPARED_BY", ""))

        sig_line = doc.add_paragraph()
        sig_line.add_run("Signature: " + "_" * 50)

        date_line = doc.add_paragraph()
        date_line.add_run("Date: " + "_" * 30)

        place_line = doc.add_paragraph()
        place_line.add_run("Place: " + "_" * 50)

        # Reviewed by
        doc.add_paragraph()
        reviewed_para = doc.add_paragraph()
        reviewed_para.add_run("Reviewed by: ").bold = True
        reviewed_para.add_run(self.change_reasons.get("REVIEWED_BY", ""))

        sig_line = doc.add_paragraph()
        sig_line.add_run("Signature: " + "_" * 50)

        date_line = doc.add_paragraph()
        date_line.add_run("Date: " + "_" * 30)

        place_line = doc.add_paragraph()
        place_line.add_run("Place: " + "_" * 50)

        # Approved by
        doc.add_paragraph()
        approved_para = doc.add_paragraph()
        approved_para.add_run("Approved by: ").bold = True
        approved_para.add_run(self.change_reasons.get("APPROVED_BY", ""))

        sig_line = doc.add_paragraph()
        sig_line.add_run("Signature: " + "_" * 50)

        date_line = doc.add_paragraph()
        date_line.add_run("Date: " + "_" * 30)

        place_line = doc.add_paragraph()
        place_line.add_run("Place: " + "_" * 50)

        # Save
        doc.save(output_path)
        print(f"\n✅ COC Report saved: {output_path}")
        return output_path


def main():
    workspace = Path(__file__).parent
    logger.debug(f"main: workspace={workspace}")
    logger.debug(f"main: DEBUG_MODE={DEBUG_MODE}")

    # Check if user wants to configure files manually
    use_file_selector = "--configure" in sys.argv or "-c" in sys.argv

    # Load existing config
    file_config = load_file_config(workspace)

    # Show file selector if requested or no config exists
    if use_file_selector or not file_config:
        print("\n" + "=" * 60)
        print("FILE SELECTION")
        print("=" * 60)
        result = show_file_selector_dialog(workspace)

        if result == "auto":
            print("User chose auto-detection")
            file_config = None
        elif result:
            file_config = result
            print("✓ Files configured successfully")
        else:
            print("❌ File selection cancelled")
            return

    # STEP 1: Extract version & variant from BOM files
    print("\n" + "=" * 60)
    print("STEP 1: EXTRACTING VERSION & VARIANT FROM BOM FILES")
    print("=" * 60)
    logger.debug("main: Calling auto_detect_config")

    if file_config:
        # Use configured files
        print(f"Using configured files from: input_files_config.json")
        logger.debug(f"Using file_config: {file_config}")

        # Set CONFIG from file paths
        CONFIG["bom_old_filename"] = Path(file_config["bom_old"]).name
        CONFIG["bom_new_filename"] = Path(file_config["bom_new"]).name
        CONFIG["schematic_old_filename"] = Path(file_config["schematic_old"]).name
        CONFIG["schematic_new_filename"] = Path(file_config["schematic_new"]).name
        CONFIG["assembly_old_filename"] = Path(file_config["assembly_old"]).name
        CONFIG["assembly_new_filename"] = Path(file_config["assembly_new"]).name

        # Extract metadata from configured BOMs
        bom_old_path = Path(file_config["bom_old"])
        bom_new_path = Path(file_config["bom_new"])

        metadata_old = extract_bom_metadata(bom_old_path)
        metadata_new = extract_bom_metadata(bom_new_path)

        # Set product info
        CONFIG["product_name"] = (
            metadata_old["product_name"] or metadata_old["assembly_no"][:7]
            if metadata_old["assembly_no"]
            else "Unknown"
        )

        CONFIG["version_old"] = (
            f"V{metadata_old['version']}" if metadata_old["version"] else "V1"
        )
        CONFIG["version_new"] = (
            f"V{metadata_new['version']}" if metadata_new["version"] else "V2"
        )

        CONFIG["variant_old"] = metadata_old["variant"] or ""
        CONFIG["variant_letter_old"] = metadata_old["variant_letter"] or ""
        CONFIG["variant_new"] = metadata_new["variant"] or ""
        CONFIG["variant_letter_new"] = metadata_new["variant_letter"] or ""

    else:
        # Auto-detect from directory
        auto_detect_config(workspace)

    # Validate that version and variant information was found
    if not CONFIG.get("bom_old_filename") or not CONFIG.get("bom_new_filename"):
        print("\n❌ ERROR: Could not find BOM files for version detection!")
        print(f"   BOM old filename: {CONFIG.get('bom_old_filename')}")
        print(f"   BOM new filename: {CONFIG.get('bom_new_filename')}")
        print("   Please run with --configure flag to manually select files:")
        print("   python create_report.py --configure")
        return

    if not CONFIG.get("version_old") or not CONFIG.get("version_new"):
        print("\n⚠️  WARNING: Version information not found in BOM metadata")
        print("   Using fallback version detection from filenames")

    # Display extracted configuration
    print("\n" + "=" * 60)
    print("AUTO-DETECTED CONFIGURATION (from BOM contents)")
    print("=" * 60)
    print(f"Product Name: {CONFIG['product_name']}")
    print(f"Old Version: {CONFIG['version_old']}")
    print(f"New Version: {CONFIG['version_new']}")
    if CONFIG.get("variant_old"):
        print(
            f"Old Variant: {CONFIG['variant_old']} {CONFIG.get('variant_letter_old', '')}"
        )
    if CONFIG.get("variant_new"):
        print(
            f"New Variant: {CONFIG['variant_new']} {CONFIG.get('variant_letter_new', '')}"
        )
    print(f"BOM Old: {CONFIG['bom_old_filename']}")
    print(f"BOM New: {CONFIG['bom_new_filename']}")
    print(f"Schematic Old: {CONFIG['schematic_old_filename']}")
    print(f"Schematic New: {CONFIG['schematic_new_filename']}")
    print(f"Assembly Old: {CONFIG['assembly_old_filename']}")
    print(f"Assembly New: {CONFIG['assembly_new_filename']}")
    print("=" * 60)

    print("\n" + "=" * 60)
    print(f"{CONFIG['product_name']} COC Report Generator")
    print("=" * 60)

    # Initialize generator
    generator = COCReportGenerator(workspace)

    # Define file paths - use configured paths if available
    if file_config:
        bom_old = Path(file_config["bom_old"])
        bom_new = Path(file_config["bom_new"])
        schematic_old = Path(file_config["schematic_old"])
        schematic_new = Path(file_config["schematic_new"])
        assembly_old = Path(file_config["assembly_old"])
        assembly_new = Path(file_config["assembly_new"])
    else:
        # Check for input_files directory
        input_dir = workspace / "input_files"
        search_dir = input_dir if input_dir.exists() else workspace

        # Define file paths from CONFIG
        bom_old = search_dir / CONFIG["bom_old_filename"]
        bom_new = search_dir / CONFIG["bom_new_filename"]
        schematic_old = search_dir / CONFIG["schematic_old_filename"]
        schematic_new = search_dir / CONFIG["schematic_new_filename"]
        assembly_old = search_dir / CONFIG["assembly_old_filename"]
        assembly_new = search_dir / CONFIG["assembly_new_filename"]

    # Store PDF paths for linking
    generator.pdf_paths = {
        "schematic_v8": schematic_old,
        "schematic_v10": schematic_new,
        "assembly_v8": assembly_old,
        "assembly_v10": assembly_new,
    }

    # Compare BOMs
    generator.compare_boms(bom_old, bom_new)

    # Analyze PDFs
    generator.analyze_pdf(schematic_old, f"Schematic {CONFIG['version_old']}")
    generator.analyze_pdf(schematic_new, f"Schematic {CONFIG['version_new']}")

    # Analyze assembly drawings in detail
    if assembly_old.exists():
        generator.changes["assembly_v8_analysis"] = generator.analyze_assembly_drawing(
            assembly_old, CONFIG["version_old"]
        )
    if assembly_new.exists():
        generator.changes["assembly_v10_analysis"] = generator.analyze_assembly_drawing(
            assembly_new, CONFIG["version_new"]
        )

    # Compare assembly drawings
    generator.compare_assembly_drawings()

    # Prompt for change reasons (questionnaire)
    generator.prompt_for_change_reasons()

    # Create reports directory if it doesn't exist
    reports_dir = workspace / "coc_reports"
    reports_dir.mkdir(exist_ok=True)

    # Generate report in the reports directory
    output_file = (
        reports_dir
        / f"COC_Report_{CONFIG['product_name']}_{CONFIG['version_old']}_to_{CONFIG['version_new']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )
    generator.generate_word_report(output_file)

    print("\n" + "=" * 60)
    print("✅ COMPLETE")
    print("=" * 60)


if __name__ == "__main__":
    main()
